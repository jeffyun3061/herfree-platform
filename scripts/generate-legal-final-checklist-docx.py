"""Herfree 실서비스 전 최종 법률 검토 체크리스트 Word 문서 생성."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "legal-final-prelaunch-checklist.docx"


def set_cell_shading(cell, fill: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_answer_line(doc: Document):
    p = doc.add_paragraph()
    p.add_run("변호사 답변:  ☐ 가능   ☐ 불가   ☐ 조건부   ☐ 수정 필요\n")
    p.add_run("조치/수정안: _________________________________________________")


def add_section_table(doc: Document, title: str, intro: str | None, rows: list[tuple[str, str, str]]):
    doc.add_heading(title, level=1)
    if intro:
        doc.add_paragraph(intro)
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Table Grid"
    headers = ("#", "현재 사실 (Herfree)", "변호사에게 받을 확답")
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        set_cell_shading(hdr[i], "E8F0EE")
    for idx, (num, fact, question) in enumerate(rows, start=1):
        row = table.rows[idx].cells
        row[0].text = num
        row[1].text = fact
        row[2].text = question
    doc.add_paragraph()


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    title = doc.add_heading("Herfree 실서비스 전 최종 법률 검토 체크리스트", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "서비스: Herfree(헤르프리) — 헤르페스(HSV) 관련 익명 건강 커뮤니티 + 본인만 보는 건강 일지 (웹)\n"
        "하지 않는 것: 진료·처방·원격진료·결제·앱·AI 의료상담·영상 직접 업로드\n"
        "인프라: AWS 한국(서울) RDS·S3·EC2·Amplify\n"
        "문의: herpfree3@gmail.com\n"
        "작성: 기술팀 사실 정리 (법률 자문 아님) · 2026년 7월"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    usage = doc.add_paragraph()
    usage.add_run("사용 방법: ").bold = True
    usage.add_run(
        "각 항목에 대해 변호사님께 ☐ 가능 / ☐ 불가 / ☐ 조건부 / ☐ 수정 필요 로 답변받고, "
        "조건부·수정 시 구체 조치를 기록합니다. 본 문서는 실제 코드·화면·약관 기준 사실 정리입니다."
    )

    add_section_table(
        doc,
        "A. 서비스 전체 · 배포 GO/NO-GO",
        None,
        [
            ("A-1", "의료행위·진단·처방 없이 「참고용 커뮤니티 + 개인 기록」만 제공", "이 서비스 형태 자체가 법적으로 운영 가능한가?"),
            ("A-2", "약관: 「의료행위 제공 안 함, 참고용」 명시", "위 문구로 의료행위·의료기기 해당 회피에 충분한가?"),
            ("A-3", "결제·유료 없음, 사업자등록 없음(또는 예정)", "사업자등록 없이 회원 모집·운영 가능한가? 최소 고지는?"),
            ("A-4", "처리방침·약관에 법인명·대표·주소·보호책임자 미기재 (이메일만)", "오픈 전 필수 기재 항목과 개인 명의 운영 가능 여부"),
            ("A-5", "—", "현재 그대로 herpfree.co.kr 공개 오픈: 가능 / 불가 / 조건부"),
            ("A-6", "—", "반드시 수정·삭제해야 할 기능 목록 (우선순위 P0/P1/P2)"),
            ("A-7", "—", "서면 의견서·수정안 제공 가능 여부"),
        ],
    )
    add_answer_line(doc)

    add_section_table(
        doc,
        "B. 회원가입 · 연령 · 인증",
        None,
        [
            ("B-1", "이메일 가입: 이메일, 비밀번호, 닉네임 (실명·주민번호 없음)", "수집 항목 최소수집 원칙에 부합하는가?"),
            ("B-2", "소셜 로그인: 카카오·네이버·구글 — 제공자 고유 ID만 저장", "OAuth 위탁·제3자 제공 고지 필요 내용"),
            ("B-3", "가입 필수 동의 4: 이용약관, 처리방침, 민감정보(건강정보), 만 14세 이상", "동의 분리·필수 구조 적법한가?"),
            ("B-4", "가입 선택 동의 2: 건강정보 통계 활용, 마케팅 수신", "선택 동의 구조·문구 적절한가?"),
            ("B-5", "14세: 체크박스만, 나이·본인인증·법정대리인 동의 없음", "충분한가? 부족 시 최소 조치(생년월일/본인인증/차단)"),
            ("B-6", "—", "성인(만 19세)만 가입해야 하는가, 14세 이상으로 충분한가?"),
            ("B-7", "일지·글쓰기는 로그인 필수", "회원가입 없이 운영(비회원)이 법적으로 더 유리한가?"),
            ("B-8", "동의 이력 DB 저장 (약관·방침 버전, 민감정보·14세·마케팅·통계)", "동의 증명·보관 기간·방법 적절한가?"),
        ],
    )
    add_answer_line(doc)

    add_section_table(
        doc,
        "C. 개인정보 · 처리방침 · 이용약관",
        None,
        [
            ("C-1", "처리방침·약관 AI 초안 → 변호사 검토 예정", "전면 재작성 vs 부분 수정 중 어느 수준?"),
            ("C-2", "수집: 이메일, 비밀번호(암호화), 닉네임, 게시글·댓글, 일지, 접속·행동 로그", "수집 항목·목적·보유 기재 충족 여부"),
            ("C-3", "AWS 서울 RDS·S3·EC2 — 회원 데이터 저장", "국외 이전 해당 없음 확인, 위탁(재위탁) 고지 문구"),
            ("C-4", "SMTP 이메일 — 비밀번호 재설정 링크 발송", "위탁·수집 항목 고지 필요 여부"),
            ("C-5", "Gabia — DNS만, 회원 데이터 저장 없음", "고지 필요/불필요"),
            ("C-6", "분석: 행동 종류만, IP·UA 해시 (원문 아님)", "개인정보 해당 여부·고지·동의 필요 여부"),
            ("C-7", "제3자 제공: 원칙 없음, 법령 요청 시 예외", "문구 적절한가"),
            ("C-8", "보유·파기: 「필요한 기간」— 구체 일수 없음", "권장 보관 기간(계정·로그·백업·동의이력·신고)"),
            ("C-9", "문의: Gmail만, 보호책임자 성명·직책 없음", "필수 기재 항목·예시 문구"),
            ("C-10", "약관 「회사」 표현 — 실체 미확정", "개인·예비법인·법인 중 어떤 표현을 써야 하나"),
        ],
    )
    add_answer_line(doc)

    add_section_table(
        doc,
        "D. 민감정보(건강정보) · 개인 일지",
        "일지는 본인만 API로 조회·수정·삭제 가능. 운영자 개별 일지 원문 조회 기능 없음(건수·익명 집계만).",
        [
            ("D-1", "항목: 날짜, 증상 유무, 강도 1~5, 전조·유발, 투약, 수면, 스트레스, 기분, 영양제·운동, 자유 메모", "수집 가능 vs 수집 자제 권장 항목"),
            ("D-2", "본인만 CRUD, 타인 접근 차단", "접근통제로 충분한가?"),
            ("D-3", "운영자: 개별 일지 원문 조회 없음", "충분한가?"),
            ("D-4", "처리방침 「운영자 필요 범위 접근」 ↔ 실제 원문 조회 없음", "문구 수정안 (운영자 접근 범위 명확화)"),
            ("D-5", "AWS 서울 DB, HTTPS, 비밀번호 암호화, 메모 DB 평문", "서버 저장 가능한가? 필드 암호화 필수 여부"),
            ("D-6", "—", "로컬(기기)만 저장으로 바꿔야 하는가?"),
            ("D-7", "민감정보 별도 필수 동의, 거부 시 가입·일지 불가", "별도 동의만으로 충분한가? 구조 변경 필요?"),
            ("D-8", "탈퇴 시 일지 전부 삭제", "법정 보관·파기 기준 적절한가?"),
        ],
    )
    add_answer_line(doc)

    add_section_table(
        doc,
        "E. 건강정보 통계 · 2차 활용",
        "※ 일지 전체를 익명화해 통계에 쓰는 구조가 아님. 선택 동의 회원의 일부 선택값만 집계.",
        [
            ("E-1", "선택 동의 회원만, 거부해도 일지·커뮤니티 이용 가능", "2차 목적(통계) 동의 구조 적법한가?"),
            ("E-2", "포함: 증상 유무, 전조·유발·투약·수면·스트레스 선택값", "포함 범위 적절한가"),
            ("E-3", "제외: 메모, 게시글·댓글, 이메일, 닉네임", "제외 범위 충분한가"),
            ("E-4", "동의 20명 미만 통계 비공개, 항목별 5명 미만 셀 숨김", "소수 추론 방지 기준 적절한가"),
            ("E-5", "마이페이지 동의 철회 → 이후 집계 제외", "철회 효과 고지·운영 적법한가"),
            ("E-6", "연구·AI학습·판매·외부 제공 포함 안 함 (방침 명시)", "현재 고지 범위 적절한가"),
            ("E-7", "—", "위 통계 구조로 가능/불가 확답"),
            ("E-8", "—", "향후 연구·제휴·AI 시 별도 동의·심의 필요 여부"),
        ],
    )
    add_answer_line(doc)

    board_intro = (
        "공개: 공지, 자유, 질문, 포비아, 증상, 확진 경험, 연애/고지, 위로/응원, 제품/루틴 후기, 비밀사연\n"
        "비공개: 운영 문의, 1:1 비밀상담(작성자+운영자), 비밀사연(운영자 열람)\n"
        "글: 제목·본문·댓글, 닉네임/익명, 사진(JPG/PNG/WEBP 10MB). 진단명·의료기록 자제 안내만(자동 필터 없음)"
    )
    add_section_table(
        doc,
        "F. 커뮤니티 · 게시판",
        board_intro,
        [
            ("F-1", "질병(헤르페스) 관련 카테고리별 게시판", "카테고리 구성 자체 가능한가?"),
            ("F-2", "이용자 증상·경험 자유 작성 (검열·필터 없음)", "허용 범위·운영자 책임"),
            ("F-3", "제품/루틴 후기 — 약·영양제·루틴 언급 가능", "의료·표시광고·약사법 리스크"),
            ("F-4", "댓글 치료법·약 추천 가능", "삭제·면책·고지 의무 범위"),
            ("F-5", "신고·숨김·삭제 (관리자)", "운영 정책·삭제 기준 필수 항목"),
            ("F-6", "비공개 게시판 운영자 전문 열람", "고지·목적·보관 어떻게 적을지"),
            ("F-7", "비밀사연 — 운영자만 열람, 작성자 목록 제목 마스킹", "적법한가, 동의·고지 추가 필요?"),
            ("F-8", "탈퇴: 게시글·댓글 익명 유지, 사진 S3 삭제", "익명 유지 vs 전 삭제 어느 쪽 적절?"),
            ("F-9", "익명 표시 vs 내부 작성자 ID 유지(신고용)", "고지 필요 여부"),
        ],
    )
    add_answer_line(doc)

    add_section_table(
        doc,
        "G. 정보 콘텐츠 (칼럼 · 유튜브)",
        None,
        [
            ("G-1", "칼럼 — 운영자 등록 일반 정보 글", "의료광고·의료법 해당 여부"),
            ("G-2", "유튜브 — 외부 영상 링크/임베드", "주의 표현·면책·출처 필수 여부"),
            ("G-3", "약관 「참고용, 의료 전문가 상담 권고」", "추가 고지(면책·응급 안내) 필요 여부"),
            ("G-4", "—", "「전문가」 표현 사용 시 자격·광고 문제"),
        ],
    )
    add_answer_line(doc)

    add_section_table(
        doc,
        "H. 외부 연동 · 상담",
        None,
        [
            ("H-1", "카카오 오픈채팅 링크 (외부 상담·문의)", "개인정보 이전·책임 (Herfree 밖) 고지"),
            ("H-2", "OAuth — 카카오·네이버·구글", "처리방침 위탁·국외(구글) 고지"),
            ("H-3", "staging·production 동일 OAuth 사용", "고지·리스크"),
        ],
    )
    add_answer_line(doc)

    add_section_table(
        doc,
        "I. 회원 탈퇴 · 권리",
        None,
        [
            ("I-1", "일지: 전부 삭제", "적법한가? 추가 보관 필요?"),
            ("I-2", "계정: 이메일·비밀번호·OAuth 비식별·삭제", "적법한가?"),
            ("I-3", "게시글·댓글: 익명 처리, 내용 유지", "허용 vs 전 삭제 권고"),
            ("I-4", "게시글 사진: S3 삭제", "적법한가?"),
            ("I-5", "마이페이지: 닉네임 수정, 탈퇴", "열람·정정·삭제·동의철회 절차 충분한가?"),
            ("I-6", "—", "재가입·동일 이메일 정책 권고"),
        ],
    )
    add_answer_line(doc)

    add_section_table(
        doc,
        "J. 운영 · 보안 · 데이터 관리",
        None,
        [
            ("J-1", "staging·production DB·S3 분리", "테스트 데이터 고지 필요?"),
            ("J-2", "RDS 백업(단기), 로그 CloudWatch", "백업·로그 보관·파기 권장 기간"),
            ("J-3", "장기 미접속 회원 — 별도 정책·자동 파기 없음", "휴면·파기 의무·권장"),
            ("J-4", "유출 시 신고·통지 — 운영 절차 확정 전", "신고·통지·기록 법적 의무"),
            ("J-5", "관리자 MFA 없음, 일지 메모 DB 평문", "안전조치(개보법 29조) 오픈 전 필수 보완 항목"),
            ("J-6", "—", "개인정보 영향평가(PIA) 필요 여부"),
            ("J-7", "—", "운영자·보호책임자 법정 의무 체크리스트"),
        ],
    )
    add_answer_line(doc)

    add_section_table(
        doc,
        "K. 금지·MVP 범위 (하지 않는 것 확인)",
        None,
        [
            ("K-1", "결제·유료·멤버십 없음", "통신판매·전자상거래 신고 불요 확인"),
            ("K-2", "네이티브 앱 없음 (웹만)", "앱스토어·의료앱 규제 해당 없음 확인"),
            ("K-3", "AI 의료 챗·자동 진단 없음", "의료기기·AI 의료 해당 없음 확인"),
            ("K-4", "영상 직접 업로드 없음 (유튜브 링크만)", "저작권·의료 콘텐츠 추가 의무"),
        ],
    )
    add_answer_line(doc)

    doc.add_heading("L. 최종 산출물 (변호사에게 요청)", level=1)
    for item in [
        "L-1  배포 가능 여부: 가능 / 불가 / 조건부 (조건 나열)",
        "L-2  P0 — 오픈 전 반드시 (기능·문구·고지·연령·운영자 정보)",
        "L-3  P1 — 오픈 후 1~3개월 (암호화·MFA·백업 등)",
        "L-4  P2 — 권장",
        "L-5  처리방침·약관·동의문 수정안 (위탁·민감정보·통계·비공개게시판 필수 조항)",
        "L-6  14세·성인·본인인증 — 최종 권고 1안",
        "L-7  서면 의견 (메일·PDF) 가능 여부",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("L 섹션 종합 답변:\n").bold = True
    p.add_run("_" * 70 + "\n" * 4)

    doc.add_heading("M. 함께 제출할 자료", level=1)
    materials = [
        "legal-client-questions-for-lawyer.pdf 또는 legal-consultation-brief.pdf",
        "이용약관·개인정보처리방침 전문 (캡처 또는 URL)",
        "회원가입 동의 6항 화면 캡처",
        "개인일지 입력 화면 캡처",
        "커뮤니티 공개·비공개 게시판 캡처",
        "마이페이지 통계 동의/철회 캡처",
        "관리자 신고·숨김 화면 캡처",
        "칼럼·유튜브 정보 화면 캡처",
        "운영 주체 (사업자/개인, 대표, 주소, 보호책임자) — 비어 있으면 「기재 예정」",
        "AWS: 한국 서울 RDS·S3·EC2, DB 직접 공개 없음, HTTPS (1문장)",
    ]
    for m in materials:
        doc.add_paragraph(f"☐ {m}", style="List Bullet")

    doc.add_paragraph()
    doc.add_paragraph("제출 금지: 회원 DB, secret, .env, 실제 일지·게시글 원문", style="Intense Quote")

    doc.add_heading("N. 운영 주체 정보 (변호사 전 기입)", level=1)
    operator_table = doc.add_table(rows=5, cols=2)
    operator_table.style = "Table Grid"
    fields = [
        ("사업자(법인)명 또는 개인 운영", ""),
        ("대표자 / 운영자", ""),
        ("사업장 주소", ""),
        ("개인정보 보호책임자 (성명·직책·연락처)", ""),
        ("사업자등록 여부", "☐ 없음  ☐ 있음  ☐ 예정"),
    ]
    for i, (label, val) in enumerate(fields):
        operator_table.rows[i].cells[0].text = label
        operator_table.rows[i].cells[1].text = val

    doc.add_heading("상담 마무리 요청 문장", level=1)
    closing = doc.add_paragraph()
    closing.add_run(
        "「위 체크리스트 A~L 전 항목에 대해 가능/불가/조건부와 P0 수정 목록을 우선순위대로 알려 주시고, "
        "처리방침·민감정보·통계·비공개게시판·14세·운영자 표기는 수정 문구 또는 재작성 필요 여부를 확답해 주세요.」"
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("\n\nHerfree Platform · 실서비스 전 법률 검토 체크리스트 · 기술팀 작성 · 법적 효력 없음")
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.save(OUT)
    print(f"Created: {OUT}")


if __name__ == "__main__":
    main()
